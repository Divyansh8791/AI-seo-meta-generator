# Downloading "en_core_web_sm" explicitly .

import spacy
import subprocess
import importlib.util

# Automatically download en_core_web_sm if not present
if not importlib.util.find_spec("en_core_web_sm"):
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])


# Extracting the text from input : 

import requests
from bs4 import BeautifulSoup
import fitz
from docx import Document

def extract_text_from_url(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        paragraphs = soup.find_all('p')
        return "\n".join(p.text for p in paragraphs)
    except Exception as e:
        return f"Error extracting from URL: {e}"

def extract_text_from_pdf(path):
    try:
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error extracting from PDF: {e}"

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting from DOCX: {e}"

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        return f"Error extracting from TXT: {e}"


# Main function to route based on input type
def extract_text(input_text=None, url=None, uploaded_file=None, file_type=None):
    try:
        if input_text:
            return input_text
        elif url:
            response = requests.get(url)
            soup = BeautifulSoup(response.text, 'html.parser')
            paragraphs = soup.find_all('p')
            return "\n".join(p.text for p in paragraphs)
        elif uploaded_file and file_type:
            if file_type == 'pdf':
                doc = fitz.open(uploaded_file)
                return "".join([page.get_text() for page in doc])
            elif file_type == 'docx':
                doc = Document(uploaded_file)
                return "\n".join([para.text for para in doc.paragraphs])
            elif file_type == 'txt':
                return uploaded_file.read().decode("utf-8")
    except Exception as e:
        return f"Error extracting text: {e}"
    return "No valid input provided."


# Extracting the keywords and topics using Spacy .

import spacy
from collections import Counter
import re

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_keywords_and_topics_spacy(text, top_k=10):
    text = clean_text(text)
    doc = nlp(text)

    # extracting keywords .
    keywords = []
    for token in doc:
        if (
            token.pos_ in ["NOUN", "PROPN"]
            and not token.is_stop
            and token.is_alpha
            and len(token.text) > 2
        ):
            keywords.append(token.lemma_.lower())
    top_keywords = Counter(keywords).most_common(top_k)

    # extracting topics .
    allowed_labels = {"ORG", "PRODUCT", "GPE", "NORP", "EVENT", "PERSON"}
    raw_topics = []
    seen = set()

    for ent in doc.ents:
        cleaned = ent.text.strip()
        if (
            ent.label_ in allowed_labels
            and len(cleaned) > 2
            and cleaned.lower() not in seen
        ):
            raw_topics.append(cleaned)
            seen.add(cleaned.lower())

    # cleaning the topics .
    filtered_topics = []
    for topic in raw_topics:
        topic_lower = topic.lower()
        if not any(
            topic_lower != other.lower() and topic_lower in other.lower()
            for other in raw_topics
        ):
            filtered_topics.append(topic)

    # final output
    cleaned_sorted_topics = sorted(set(filtered_topics))
    cleaned = []
    for t in cleaned_sorted_topics:
      t_clean = re.sub(r"^the\s+", "", t, flags=re.I).strip()  # Remove leading 'the'
      cleaned.append(t_clean)
    cleaned_sorted_topics = sorted(set(cleaned))
    return top_keywords, cleaned_sorted_topics


# Generating SEO Meta Title and Meta Description using LLM 

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain.chains import LLMChain
import json
import os


api_key = os.environ.get("GEMINI_API_KEY")
llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        api_key=api_key,  # replace with HF secret key
        temperature=0.2
    )
# Prompt for SEO title & description
seo_prompt = PromptTemplate(
    input_variables=["keywords", "topics"],
    template="""
You are an expert SEO assistant.

Using the following keywords and topics extracted from a blog post, your task is to generate:
1. A compelling SEO meta title (max 60 characters) that is clear, specific, and attractive to readers.
2. A concise SEO meta description (max 160 characters) that summarizes the blog post and uses important keywords and topics naturally.

Do not make up unrelated content. Just use the provided keywords and topics.

Respond in this JSON format:
{{
  "meta_title": "Your title here",
  "meta_description": "Your description here"
}}
but without : ``` ```

Keywords: {keywords}
Topics: {topics}
"""
)

seo_chain = LLMChain(llm=llm, prompt=seo_prompt)

# Final function
def generate_seo_metadata(keywords: list, topics: list) -> dict:
    response = seo_chain.invoke({
        "keywords": json.dumps(keywords),
        "topics": json.dumps(topics)
    })
    try:
        return json.loads(response["text"].strip("` \n"))
    except json.JSONDecodeError:
        return {"meta_title": "", "meta_description": ""}

# Final Gradio app :

import gradio as gr
import pandas as pd

# Logics for Gradio
def pipeline(input_text, url, file, file_type):
    content = extract_text(input_text=input_text, url=url, uploaded_file=file, file_type=file_type)
    keywords, topics = extract_keywords_and_topics_spacy(content)
    return content, keywords, topics

def seo_pipeline(keywords_df, topics_df):
    
    if isinstance(keywords_df, pd.DataFrame):
        keywords = [(row[0], row[1]) for row in keywords_df.values.tolist()]
    else:
        keywords = keywords_df

    if isinstance(topics_df, pd.DataFrame):
        topics = topics_df.iloc[:, 0].tolist() if not topics_df.empty else []
    else:
        topics = topics_df

    seo_result = generate_seo_metadata(keywords, topics)
    return seo_result["meta_title"], seo_result["meta_description"]

import os
import pandas as pd
from datetime import datetime

def download_csv(keywords_df, topics_df, title, description):
    # Convert keywords
    if isinstance(keywords_df, pd.DataFrame):
        keywords = [row[0] for row in keywords_df.values.tolist()]
    else:
        keywords = [kw[0] for kw in keywords_df]

    # Convert topics
    if isinstance(topics_df, pd.DataFrame):
        topics = topics_df.iloc[:, 0].tolist() if not topics_df.empty else []
    else:
        topics = topics_df

    # Create DataFrame
    df = pd.DataFrame({
        "Keywords": [", ".join(keywords)],
        "Topics": [", ".join(topics)],
        "Meta Title": [title],
        "Meta Description": [description]
    })

    #Handle output path
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"seo_output_{timestamp}.csv"

    # On Hugging Face, /mnt/data works. Else use local dir.
    output_dir = "/mnt/data" if os.path.exists("/mnt/data") else "."
    path = os.path.join(output_dir, filename)

    # Save CSV
    df.to_csv(path, index=False)
    return path

# UI for users 

with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("## ✨ AI-Powered SEO Assistant")
    gr.Markdown("Extract keywords, topics, SEO meta title & description from a blog post (Text, URL, PDF, DOCX, or TXT).")

    with gr.Row():
        input_text = gr.Textbox(label="📝 Paste Text", lines=6, placeholder="Or leave empty if uploading a file or using URL")
        url = gr.Textbox(label="🔗 URL")

    with gr.Row():
        file = gr.File(label="📁 Upload File", file_types=[".pdf", ".docx", ".txt"])
        file_type = gr.Radio(["pdf", "docx", "txt"], label="File Type")

    extract_btn = gr.Button("🚀 Extract Keywords & Topics")

    content = gr.Textbox(label="📄 Extracted Content", lines=10, interactive=False)
    keywords_output = gr.Dataframe(label="🔑 Top Keywords", headers=["Keyword", "Frequency"])
    topics_output = gr.Dataframe(label="📚 Topics")

    seo_btn = gr.Button("🎯 Generate SEO Title & Description")
    title_out = gr.Textbox(label="📌 SEO Meta Title")
    desc_out = gr.Textbox(label="📝 SEO Meta Description")

    download_btn = gr.Button("📥 Download CSV")
    download_file = gr.File(label="Download Link")

    extract_btn.click(pipeline,
                      inputs=[input_text, url, file, file_type],
                      outputs=[content, keywords_output, topics_output])

    seo_btn.click(seo_pipeline,
                  inputs=[keywords_output, topics_output],
                  outputs=[title_out, desc_out])

    download_btn.click(download_csv,
                       inputs=[keywords_output, topics_output, title_out, desc_out],
                       outputs=download_file)

demo.launch(share=True)
